import os
from pathlib import Path
from docx import Document
import pandas as pd
from typing import List, Dict, Any, Tuple


def load_docx(file_path: Path) -> List[Dict[str, Any]]:
    doc = Document(file_path)
    content_items = []
    
    current_section = "前言"
    current_level = 0
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        
        level = 0
        if para.style.name.startswith('Heading'):
            try:
                level = int(para.style.name.split(' ')[-1])
            except (IndexError, ValueError):
                level = 1
            current_section = text
            current_level = level
        
        content_items.append({
            "type": "paragraph",
            "text": text,
            "section": current_section,
            "level": level,
            "page": None
        })
    
    for table_idx, table in enumerate(doc.tables):
        table_text = []
        for row in table.rows:
            row_text = " | ".join([cell.text.strip() for cell in row.cells])
            table_text.append(row_text)
        
        if table_text:
            content_items.append({
                "type": "table",
                "text": "\n".join(table_text),
                "section": current_section,
                "level": current_level,
                "page": None
            })
    
    return content_items


def load_xlsx(file_path: Path) -> List[Dict[str, Any]]:
    content_items = []
    filename = os.path.basename(file_path)
    standard_type = get_standard_type(filename)
    
    for engine in ["calamine", "openpyxl"]:
        try:
            excel = pd.ExcelFile(file_path, engine=engine)
            break
        except Exception as e:
            if engine == "calamine":
                print(f"calamine engine failed for {filename}, trying openpyxl: {e}")
            else:
                print(f"openpyxl engine also failed for {filename}: {e}")
                return []
    
    for sheet_name in excel.sheet_names:
        try:
            df = pd.read_excel(file_path, sheet_name=sheet_name, engine=engine, dtype=str)
            df = df.fillna("")
            
            for idx, row in df.iterrows():
                row_text = f"Sheet: {sheet_name}\nRow: {idx + 1}"
                has_content = False
                
                for col in df.columns:
                    val = str(row[col]).strip()
                    if val:
                        row_text += f"\n{col}: {val}"
                        has_content = True
                
                if has_content:
                    content_items.append({
                        "text": row_text,
                        "source_file": filename,
                        "standard_type": standard_type,
                        "section": sheet_name,
                        "content_type": "xlsx_row",
                        "page_or_sheet": sheet_name
                    })
        except Exception as e:
            print(f"Failed to read sheet '{sheet_name}' in {filename}: {e}")
            continue
    
    return content_items


def get_standard_type(filename: str) -> str:
    filename_lower = filename.lower()
    if "rubric" in filename_lower or "rubrics" in filename_lower:
        return "rubric_quality"
    elif "题目质量" in filename:
        return "prompt_quality"
    elif "文生文" in filename or "评测标准" in filename:
        return "answer_evaluation"
    elif "字数" in filename:
        return "word_count_rule"
    elif "标签体系" in filename:
        return "label_taxonomy"
    elif "人设" in filename:
        return "persona_info"
    else:
        return "general_standard"
